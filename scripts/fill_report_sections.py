from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "Report - Luca Revilla.docx"
BACKUP = ROOT / "Report - Luca Revilla.backup_before_codex.docx"
OUT = DOCX
ALT_OUT = ROOT / "Report - Luca Revilla - Codex Filled.docx"
ANALYSIS = ROOT / "outputs" / "engineering_analysis"


def paragraph_index(doc: Document, prefix: str) -> int:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return i
    raise ValueError(f"Could not find paragraph starting with {prefix!r}")


def move_before_marker(element, marker):
    marker.addprevious(element)


def add_para(doc: Document, marker, text: str = "", style: str | None = None, bold: bool = False):
    paragraph = doc.add_paragraph(style=style)
    if text:
        run = paragraph.add_run(text)
        run.bold = bold
    move_before_marker(paragraph._p, marker)
    return paragraph


def add_bullets(doc: Document, marker, items: list[str]):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        move_before_marker(paragraph._p, marker)


def add_table(doc: Document, marker, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    move_before_marker(table._tbl, marker)
    return table


def add_figure(doc: Document, marker, image_path: Path, caption: str, width: float = 5.9):
    if not image_path.exists():
        add_para(doc, marker, f"[Missing figure: {image_path.name}]", style="Normal")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    move_before_marker(paragraph._p, marker)
    caption_p = doc.add_paragraph(style="Normal")
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    caption_run.italic = True
    caption_run.font.size = Pt(9)
    move_before_marker(caption_p._p, marker)


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    doc = Document(DOCX)
    start = paragraph_index(doc, "6. Critical Evaluation")
    end = paragraph_index(doc, "8. Conclusions")
    marker = doc.paragraphs[end]._p

    body = doc._body._element
    to_remove = [p._p for p in doc.paragraphs[start:end]]
    for element in to_remove:
        body.remove(element)

    add_para(doc, marker, "6. Critical Evaluation of the Paper", style="Heading 1")
    add_para(doc, marker, "6.1 Novelty and Significance", style="Heading 2")
    add_para(
        doc,
        marker,
        "Howell et al. make a meaningful contribution because they move beyond image-level lung ultrasound "
        "classification toward real-time, multiclass semantic segmentation of artefacts and anatomy. This is "
        "important because LUS interpretation depends on spatial artefact patterns rather than on direct tissue "
        "visualization alone. A segmentation model can therefore provide two outputs that are more useful than a "
        "single classification label: a visual overlay that may support training, and a structured mask from which "
        "quantitative measures such as the B-line Artefact Score (BLAS) can be computed.",
    )
    add_para(
        doc,
        marker,
        "The strongest aspect of the paper is this connection between real-time engineering implementation and a "
        "clinically interpretable downstream metric. The work is not only an accuracy benchmark; it attempts to "
        "translate segmentation into visual guidance and semi-quantitative severity assessment. That makes the paper "
        "well suited for engineering evaluation because the central claims can be tested quantitatively at several "
        "levels: pixel segmentation, BLAS agreement, robustness to segmentation errors, uncertainty, and deployment cost.",
    )

    add_para(doc, marker, "6.2 Dataset Adequacy and External Validity", style="Heading 2")
    add_para(
        doc,
        marker,
        "The phantom dataset is a strength and a limitation. It provides a controlled platform with labelled examples "
        "of ribs, pleural line, A-lines, B-lines, and B-line confluence, and the authors made the phantom data and code "
        "available. This substantially improves reproducibility. The use of a commercial training phantom is also "
        "reasonable for a first demonstration because it reduces biological variability and allows the authors to focus "
        "on the engineering feasibility of real-time segmentation.",
    )
    add_para(
        doc,
        marker,
        "However, external validity remains limited. Phantom images do not capture the full variability of real patients, "
        "including body habitus, probe pressure, respiratory motion, scanner presets, transducer differences, pathology "
        "heterogeneity, and operator-dependent acquisition. The clinical transfer-learning dataset was small, contained "
        "57 images from 41 patients, and lacked sufficient examples of several artefact classes. The authors state that "
        "the clinical data were fully anonymized, making patient-level splitting impossible; therefore, those clinical "
        "results should be interpreted as preliminary feasibility rather than evidence of clinical decision-support readiness.",
    )

    add_para(doc, marker, "6.3 Annotation Quality and Ground-Truth Uncertainty", style="Heading 2")
    add_para(
        doc,
        marker,
        "The annotation strategy is credible but not definitive. Images were labelled using polygon annotations for five "
        "foreground classes and converted to segmentation masks. Multiple annotators with different levels of ultrasound "
        "experience contributed labels, and a senior sonographer labelled additional images with peer review. This is a "
        "practical approach for a limited dataset, but it also introduces label uncertainty. LUS artefacts are inherently "
        "subjective: the boundary between B-line and B-line confluence, or between true artefact and background speckle, "
        "can be ambiguous.",
    )
    add_para(
        doc,
        marker,
        "The paper would be stronger with a dedicated inter- and intra-annotator variability analysis. Without it, Dice "
        "scores are difficult to interpret fully: a lower Dice for B-lines may reflect model failure, label ambiguity, "
        "or both. This is especially important because BLAS is computed from B-line and confluence masks. If the ground "
        "truth for these classes is uncertain, the downstream severity metric inherits that uncertainty.",
    )

    add_para(doc, marker, "6.4 Model Design and Training Choices", style="Heading 2")
    add_para(
        doc,
        marker,
        "The lightweight U-Net design is appropriate for the authors' real-time objective. U-Net architectures are well "
        "matched to biomedical segmentation because they combine local spatial detail with broader image context. The "
        "authors' comparison with more complex U-Net variants also supports their design choice: a smaller model that "
        "performs comparably while running faster is preferable for point-of-care ultrasound.",
    )
    add_para(
        doc,
        marker,
        "The augmentation strategy is another strength. Ultrasound-specific augmentations, including gain, time gain "
        "compensation, and depth changes, are particularly relevant because artefact appearance is sensitive to acquisition "
        "settings. Validation monitoring, early stopping, and repeated test splits are also appropriate for a small, "
        "class-imbalanced segmentation dataset. The main missing element is a stronger uncertainty and failure-mode analysis: "
        "the model is evaluated primarily by overlap metrics, but the clinical risk depends on which classes fail and how "
        "those failures propagate into BLAS.",
    )

    add_para(doc, marker, "6.5 Validation Strategy and Risk of Bias", style="Heading 2")
    add_para(
        doc,
        marker,
        "The authors explicitly attempted to reduce leakage in the phantom study by splitting at the video level, reducing "
        "the likelihood that nearly identical frames appear in both train and test sets. This is an important design choice "
        "because adjacent frames from ultrasound clips can be highly redundant. In the local reproduction, the train/test "
        "filenames did not share sequence identifiers, which is consistent with a video- or sequence-aware split.",
    )
    add_para(
        doc,
        marker,
        "Nevertheless, the phantom domain remains visually redundant. A nearest-neighbor analysis performed for this report "
        "found no shared sequence IDs between train and test, but 17 of 100 test images had a nearest train image with "
        "cosine similarity >= 0.99 after cropped low-resolution normalization, and 41 of 100 had similarity >= 0.95. "
        "This does not prove leakage, but it suggests that phantom test performance may overestimate generalization to "
        "new scanners, patients, acquisition sites, or disease phenotypes.",
    )
    add_figure(
        doc,
        marker,
        ANALYSIS / "train_test_similarity_hist.png",
        "Figure 1. Nearest-neighbor visual similarity between test images and the training set in the local reproduction.",
    )

    add_para(doc, marker, "6.6 Clinical Meaningfulness of Reported Metrics", style="Heading 2")
    add_para(
        doc,
        marker,
        "Dice and IoU are necessary but insufficient for evaluating this application. A model can achieve high pixel accuracy "
        "because background is easy, while still making clinically meaningful errors in the relatively small B-line or "
        "confluence regions. The paper's own motivation points toward a downstream biomarker, BLAS, so the central clinical "
        "question is not only whether masks overlap, but whether segmentation errors change the inferred artefact burden.",
    )
    add_para(
        doc,
        marker,
        "The clinical transfer-learning results also limit the strength of the paper's translational claim. The phantom model "
        "is technically feasible and potentially useful for education or real-time visual feedback, but clinical decision "
        "support would require prospective validation, patient-level splitting, uncertainty handling, and evidence that the "
        "derived severity metrics correspond to clinically meaningful outcomes.",
    )

    add_para(doc, marker, "6.7 Interpretability, Safety, and Deployment", style="Heading 2")
    add_para(
        doc,
        marker,
        "The segmentation overlay is an interpretable output relative to black-box classification, but interpretability does "
        "not guarantee safety. A confident-looking overlay may increase user trust even when the model is wrong. This is "
        "especially relevant for B-line confluence, where false positives can make a frame appear more severe, and false "
        "negatives can make a high-burden frame appear reassuring.",
    )
    add_para(
        doc,
        marker,
        "For deployment, the paper appropriately considers inference speed, but a complete clinical engineering analysis "
        "should also consider calibration, uncertainty, failure detection, power consumption, battery life, workflow impact, "
        "and scanner-to-scanner generalization. The model is promising as an educational or assistive visualization tool; "
        "it is not yet validated as an autonomous severity-scoring system.",
    )

    add_para(doc, marker, "6.8 Overall Strengths, Weaknesses, and Missing Experiments", style="Heading 2")
    add_bullets(
        doc,
        marker,
        [
            "Strengths: open phantom data and code, real-time implementation, ultrasound-specific augmentation, a lightweight model appropriate for PoCUS, and an interpretable segmentation-derived BLAS metric.",
            "Weaknesses: reliance on phantom images, limited clinical validation, label subjectivity, incomplete annotator-variability analysis, and uncertain generalization across patients, scanners, and acquisition protocols.",
            "Highest-priority missing experiments: patient-level prospective validation, scanner/transducer external validation, annotator variability, task-specific uncertainty estimation, and BLAS robustness to segmentation failure modes.",
        ],
    )

    add_para(doc, marker, "7. Quantitative Engineering Analysis Applied to Medicine/Biology", style="Heading 1")
    add_para(doc, marker, "7.1 Engineering Question", style="Heading 2")
    add_para(
        doc,
        marker,
        "The engineering analysis asks whether a reproduced LUS segmentation model can reliably support the downstream "
        "BLAS severity metric. The core hypothesis is that standard segmentation metrics are not sufficient: the clinically "
        "relevant risk is whether class-specific errors in B-line and confluence regions cause BLAS to under- or over-estimate "
        "vertical artefact burden.",
    )

    add_para(doc, marker, "7.2 Data Sources and Reproduction Setup", style="Heading 2")
    add_para(
        doc,
        marker,
        "The analysis used the public phantom frames available in the repository, with 464 paired training images and 100 "
        "paired test images in the local workspace. The original TensorFlow/Keras model weights (`model_lus.h5`) were loaded "
        "and evaluated on the test split using the same preprocessing convention used in the reproduction scripts: crop "
        "[100, 50, 850, 460], resize to 256 x 256 pixels, grayscale normalization, and multiclass argmax masks with labels "
        "0 through 5. Prediction masks were saved both as raw label images and as colored visualizations for review.",
    )
    add_para(
        doc,
        marker,
        "The reproduced test-set segmentation achieved pixel accuracy of 0.958 and mean Dice excluding background of 0.701. "
        "Class-wise Dice values were 0.777 for ribs, 0.791 for pleural line, 0.660 for A-line, 0.633 for B-line, and 0.645 "
        "for B-line confluence. These results were then used as the basis for downstream BLAS agreement, failure-case, "
        "sensitivity, uncertainty, and deployment analyses.",
    )

    add_para(doc, marker, "7.3 BLAS Agreement Between Manual and Predicted Masks", style="Heading 2")
    add_para(
        doc,
        marker,
        "BLAS was computed from both manual masks and predicted masks. Agreement was assessed using mean absolute error, "
        "root mean squared error, correlation, Bland-Altman bias, and category disagreement using approximate categories "
        "of low (<0.5), intermediate (0.5-0.9), and high (>0.9) BLAS.",
    )
    add_table(
        doc,
        marker,
        ["Metric", "Result", "Interpretation"],
        [
            ["N test frames", "100", "Full local test split."],
            ["BLAS MAE", "0.158", "Average absolute downstream score error."],
            ["BLAS RMSE", "0.284", "Large errors occur in a subset of frames."],
            ["Bias (predicted - manual)", "+0.060", "Predicted masks slightly overestimate BLAS on average."],
            ["Pearson r", "0.626", "Moderate linear agreement."],
            ["Spearman rho", "0.595", "Moderate rank agreement."],
            ["Category disagreement", "26/100", "One quarter of frames changed severity category."],
        ],
    )
    add_figure(
        doc,
        marker,
        ANALYSIS / "blas_scatter.png",
        "Figure 2. Manual-label BLAS versus predicted-mask BLAS for the reproduced test set.",
    )
    add_figure(
        doc,
        marker,
        ANALYSIS / "blas_bland_altman.png",
        "Figure 3. Bland-Altman analysis showing BLAS bias and limits of agreement.",
    )
    add_para(
        doc,
        marker,
        "The agreement analysis shows that the model is useful but not yet reliable enough for unqualified severity scoring. "
        "A mean absolute error of 0.158 is nontrivial on a 0-1 scale, and 26% category disagreement means that downstream "
        "clinical interpretation could change even when pixel-level segmentation metrics appear acceptable.",
    )

    add_para(doc, marker, "7.4 Failure-Case Analysis", style="Heading 2")
    add_para(
        doc,
        marker,
        "The ten largest BLAS errors were inspected visually. The dominant failure pattern was false-positive B-line "
        "confluence in images whose manual labels had low or zero BLAS. For example, S316-F39, S316-F78, and S316-F66 all "
        "had manual BLAS of 0.0 but predicted BLAS values above 0.95 because the model predicted broad confluence regions. "
        "The opposite failure also occurred: S80-F19 had manual BLAS of 0.980 but predicted BLAS of 0.267, indicating "
        "under-detection of confluence burden.",
    )
    add_figure(
        doc,
        marker,
        ANALYSIS / "failure_cases" / "01_S316-F39.png",
        "Figure 4. Example false-positive confluence failure case with high predicted BLAS despite zero manual BLAS.",
    )
    add_figure(
        doc,
        marker,
        ANALYSIS / "failure_cases" / "07_S80-F19.png",
        "Figure 5. Example under-detection failure case with high manual BLAS and low predicted BLAS.",
    )

    add_para(doc, marker, "7.5 Sensitivity Analysis of BLAS", style="Heading 2")
    add_para(
        doc,
        marker,
        "A sensitivity analysis perturbed the B-line and confluence regions to quantify how segmentation errors propagate "
        "into BLAS. False-negative perturbations removed B-line/confluence pixels, false-positive perturbations added B-line "
        "pixels within the BLAS region, and morphological erosion/dilation simulated systematic under- or over-segmentation.",
    )
    add_para(
        doc,
        marker,
        "The resulting curves show that BLAS is especially sensitive to spatial extent errors in the vertical artefact regions. "
        "Five iterations of dilation increased BLAS by a mean of 0.124, while five iterations of erosion decreased BLAS by "
        "a mean of 0.171. This supports the central engineering conclusion: BLAS is an interpretable metric, but its reliability "
        "depends strongly on robust segmentation of B-line and confluence morphology.",
    )
    add_figure(
        doc,
        marker,
        ANALYSIS / "sensitivity_curves.png",
        "Figure 6. Sensitivity of BLAS to controlled B-line/confluence false positives, false negatives, erosion, and dilation.",
    )

    add_para(doc, marker, "7.6 Conformal Prediction and Uncertainty", style="Heading 2")
    add_para(
        doc,
        marker,
        "Conformal prediction was used as a lightweight uncertainty analysis. Instead of forcing each pixel to take only the "
        "argmax class, the method constructs a set of plausible classes using calibration scores of 1 minus the probability "
        "assigned to the true class. With alpha = 0.1, the global pixel-level analysis achieved mean held-out coverage of "
        "0.905, close to the nominal 90% target.",
    )
    add_para(
        doc,
        marker,
        "However, the global conformal result is dominated by background pixels. A BLAS-focused refinement repeated the "
        "analysis on foreground pixels, the manual BLAS ROI, and manual B-line/confluence pixels. The correlation between "
        "mean entropy and absolute BLAS error remained weak: Spearman rho was 0.195 for all pixels, 0.218 for the BLAS ROI, "
        "and approximately zero for manual B-line/confluence pixels. Thus, pixel-level uncertainty alone did not reliably "
        "identify frames with large downstream BLAS error.",
    )
    add_table(
        doc,
        marker,
        ["Conformal subset", "Coverage / correlation result", "Interpretation"],
        [
            ["All pixels", "Coverage 0.905; rho 0.195", "Nominal calibration, but diluted by background."],
            ["Foreground pixels", "rho 0.017", "Foreground uncertainty alone did not explain BLAS error."],
            ["Manual BLAS ROI", "rho 0.218", "Slightly more relevant, but still weak."],
            ["Manual B-line/confluence", "rho -0.027", "Class-specific uncertainty did not track downstream error."],
        ],
    )
    add_figure(
        doc,
        marker,
        ANALYSIS / "conformal_prediction_set_cases" / "01_S316-F39_prediction_sets.png",
        "Figure 7. Conformal prediction-set visualization for a high-error case. The model's plausible set strongly includes confluence in the predicted BLAS region.",
    )

    add_para(doc, marker, "7.7 Portable Deployment and Energy Estimate", style="Heading 2")
    add_para(
        doc,
        marker,
        "The reproduced model has 7.86 million parameters and an estimated 13.9 GMAC per 256 x 256 frame, counting Conv2D "
        "and Conv2DTranspose layers. Local CPU TensorFlow inference took 0.147 seconds per frame, corresponding to 6.82 FPS. "
        "Using simple power assumptions, this corresponds to approximately 3.67 J/frame at 25 W on a laptop-class CPU, "
        "2.20 J/frame at 15 W if the same latency were achieved on a Jetson-class device, and 0.5 J/frame for a 5 W edge "
        "accelerator running at 10 FPS.",
    )
    add_table(
        doc,
        marker,
        ["Platform assumption", "FPS", "Energy/frame", "50 Wh battery runtime"],
        [
            ["Measured local CPU, 25 W", "6.82", "3.67 J", "2.0 h"],
            ["Jetson-class 15 W, same latency", "6.82", "2.20 J", "3.3 h"],
            ["5 W edge accelerator at 10 FPS", "10.0", "0.50 J", "10.0 h"],
            ["3 W portable CPU at 1 FPS", "1.0", "3.00 J", "16.7 h"],
        ],
    )
    add_para(
        doc,
        marker,
        "These estimates suggest that real-time portable inference is technically plausible, but a deployable ultrasound "
        "system would need hardware-specific benchmarking, thermal testing, and integration with the scanner display pipeline. "
        "The original paper's real-time claim is credible for model feasibility, but energy and battery behavior should be "
        "measured directly on the intended portable platform.",
    )

    add_para(doc, marker, "7.8 Train-Test Similarity and Leakage Probe", style="Heading 2")
    add_para(
        doc,
        marker,
        "Because ultrasound videos can contain many visually similar frames, a leakage probe compared each test image with "
        "its nearest training image using cropped, normalized low-resolution image features. No train/test sequence IDs were "
        "shared in the local split, which supports the authors' claim of leakage reduction. However, visual redundancy remained: "
        "17 of 100 test images had nearest-neighbor cosine similarity >= 0.99, and 41 of 100 were >= 0.95.",
    )
    add_para(
        doc,
        marker,
        "This result should be interpreted carefully. It does not demonstrate improper leakage, but it does show that the "
        "phantom domain contains repeated visual patterns across nominally separate sequences. Therefore, the reproduced "
        "test performance should be framed as controlled phantom reproducibility rather than proof of clinical generalization.",
    )

    add_para(doc, marker, "7.9 Engineering Interpretation", style="Heading 2")
    add_para(
        doc,
        marker,
        "The engineering analysis supports a nuanced conclusion. The paper demonstrates that lightweight real-time LUS "
        "segmentation is feasible and that segmentation masks can produce an interpretable BLAS metric. However, the reproduced "
        "results show that BLAS is sensitive to class-specific segmentation errors, especially false-positive or false-negative "
        "B-line confluence. Uncertainty analysis and train-test similarity probing further indicate that a clinically reliable "
        "system would require task-specific failure detection, stronger external validation, and hardware-specific deployment "
        "testing before use in triage or decision support.",
    )

    try:
        doc.save(OUT)
        print(f"Updated {OUT}")
    except PermissionError:
        doc.save(ALT_OUT)
        print(f"Original file was locked; wrote {ALT_OUT}")
    print(f"Backup at {BACKUP}")


if __name__ == "__main__":
    main()
